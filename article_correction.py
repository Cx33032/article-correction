###################################################

# @Time: Jan 30, 2024

# @Developer: Harryjin

# @Description: 通过DashScope API进行文章校对和翻译

###################################################

import argparse
from http import HTTPStatus
import dashscope
import docx
import re
import json
from tqdm import trange

dashscope.api_key = "YOUR-API-KEY" # API秘钥

"""
API-KEY 添加方法
1. 访问 https://dashscope.console.aliyun.com/apiKey
2. 点击“创建新的API-KEY”
3. 点击蓝色的“复制”按钮
4. 将其复制并替换上方 YOUR-API-KEY （切记不要删除上方的双引号）

如果未开通DashScope服务， 请查阅https://help.aliyun.com/zh/dashscope/developer-reference/activate-dashscope-and-create-an-api-key?spm=a2c4g.11186623.0.i1
需要实名阿里云Aliyun账户并且账户未欠费才可以使用
"""

def check_error(artIn):
    messages = [{'role': 'system', 'content': 'You are a professional editor who checks spelling and punctuational mistakes in the article'},
                {'role': 'user', 'content': '纠正以下英文片段的标点以及拼写错误：' + artIn}]

    response = dashscope.Generation.call(
        dashscope.Generation.Models.qwen_turbo,
        messages=messages,
        result_format='message',
    )
    if response.status_code == HTTPStatus.OK:
        return response
    else:
        print('Request id: %s, Status code: %s, error code: %s, error message: %s' % (
            response.request_id, response.status_code,
            response.code, response.message
        ))

def translate_article(artIn):
    messages = [{'role': 'system', 'content': 'You are a professional translator who mainly translates English books to Chinese which are majored in Medication'},
                {'role': 'user', 'content': '请翻译下列医学专业书籍，将英文翻译为中文，结合上下文，注重自然性：' + artIn}]

    response = dashscope.Generation.call(
        dashscope.Generation.Models.qwen_turbo,
        messages=messages,
        result_format='message',
    )
    if response.status_code == HTTPStatus.OK:
        return response
    else:
        print('Request id: %s, Status code: %s, error code: %s, error message: %s' % (
            response.request_id, response.status_code,
            response.code, response.message
        ))

def checkLetters(inStr):
   check = re.compile(r'[A-Za-z]', re.S)
   res = re.findall(check, inStr)
   if len(res):
      return True
   else:
      return False 

if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('-i', '--input', type=str, default='input.docx', help='Input docx file')
    parser.add_argument('-m', '--method', type=str, default='check', help='Correction | Translation')
    args = parser.parse_args()
    inputFile = args.input
    method = args.method
    outputFile = str(inputFile).replace('.docx', '_after.docx')

    doc = docx.Document(inputFile)
    newDoc = docx.Document()

    for i in trange(len(doc.paragraphs)):
        artIn = doc.paragraphs[i].text
        if(checkLetters(artIn)):
            if(method == 'check'):
                rawOut = check_error(artIn)
            else:
                rawOut = translate_article(artIn)
            temp = json.dumps(rawOut)
            temp = temp.replace('\\\"', '')
            artOut = json.loads(temp)
            newDoc.add_paragraph(artOut['output']['choices'][0]['message']['content'])
        else:
            newDoc.add_paragraph(artIn)
    
    newDoc.save(outputFile)
